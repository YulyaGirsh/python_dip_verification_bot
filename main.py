import asyncio
from aiogram import Bot, Dispatcher, types
from aiogram.filters import CommandStart
from aiogram.utils.keyboard import InlineKeyboardBuilder
import docx
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn
import os
import re
from typing import List, Tuple, Dict, Set
from collections import defaultdict

BOT_TOKEN = "8141487259:AAEYJ9Qtbrdaz411UP35qhY7f8OVdNFZJ_s"

# Настройки проверки
REQUIRED_HEADERS_LOWER = ["введение", "список источников литературы", "заключение"]
OPTIONAL_HEADERS_LOWER = ["содержание", "приложение"]
MAIN_FONT_SIZE = 14
MAIN_FONT_NAME = "Times New Roman"
LINE_SPACING = 1.5
PARAGRAPH_INDENT_CM = 1.25  # Отступ первой строки в см
APPENDIX_PATTERN = re.compile(r'^Приложение [А-ДЕ-Я]$', re.IGNORECASE)
FORBIDDEN_LETTERS = {'Ё', 'И'}  # Буквы, которые нельзя использовать в приложениях
APPENDIX_ORDER = ['А', 'Б', 'В', 'Г', 'Д', 'Е', 'Ж', 'З', 'Й', 'К', 'Л', 'М',
                  'Н', 'О', 'П', 'Р', 'С', 'Т', 'У', 'Ф', 'Х', 'Ц', 'Ч', 'Ш', 'Щ',
                  'Ъ', 'Ы', 'Ь', 'Э', 'Ю', 'Я']

# Требования к полям (в мм)
LEFT_MARGIN = 25
RIGHT_MARGIN = 15
TOP_MARGIN = 20
BOTTOM_MARGIN = 20
MARGIN_TOLERANCE = 1  # Допустимое отклонение в мм
PARAGRAPH_SPACING = 0  # Интервалы перед и после абзаца должны быть 0 pt

bot = Bot(token=BOT_TOKEN)
dp = Dispatcher()

# Глобальная переменная для хранения временных файлов
user_documents = {}


def is_bold(paragraph):
    """Проверяет, является ли текст жирным (учитывает и стиль, и прямое форматирование)"""
    if paragraph.style.font.bold and any(run.font.bold is False for run in paragraph.runs):
        return False
    if paragraph.style.font.bold:
        return True
    for run in paragraph.runs:
        if run.font.bold:
            return True
    return False


def is_list_paragraph(paragraph):
    """Проверяет, является ли параграф элементом списка"""
    return paragraph._element.pPr is not None and paragraph._element.pPr.numPr is not None


def get_alignment(paragraph):
    """Получает выравнивание параграфа (учитывает прямое форматирование)"""
    if paragraph.alignment is not None:
        return paragraph.alignment
    if hasattr(paragraph.style, 'paragraph_format') and paragraph.style.paragraph_format.alignment is not None:
        return paragraph.style.paragraph_format.alignment
    pPr = paragraph._element.pPr
    if pPr is not None:
        jc = pPr.find(qn('w:jc'))
        if jc is not None:
            val = jc.get(qn('w:val'))
            if val == 'center':
                return WD_ALIGN_PARAGRAPH.CENTER
            elif val == 'right':
                return WD_ALIGN_PARAGRAPH.RIGHT
            elif val == 'left':
                return WD_ALIGN_PARAGRAPH.LEFT
            elif val == 'both':
                return WD_ALIGN_PARAGRAPH.JUSTIFY
    return WD_ALIGN_PARAGRAPH.LEFT


def check_margins(section):
    """Проверяет соответствие полей требованиям"""
    errors = []
    left_mm = section.left_margin.mm
    right_mm = section.right_margin.mm
    top_mm = section.top_margin.mm
    bottom_mm = section.bottom_margin.mm

    if not (LEFT_MARGIN - MARGIN_TOLERANCE <= left_mm <= LEFT_MARGIN + MARGIN_TOLERANCE):
        errors.append(f"Левое поле: {left_mm:.1f} мм (должно быть {LEFT_MARGIN} мм)")
    if not (RIGHT_MARGIN - MARGIN_TOLERANCE <= right_mm <= RIGHT_MARGIN + MARGIN_TOLERANCE):
        errors.append(f"Правое поле: {right_mm:.1f} мм (должно быть {RIGHT_MARGIN} мм)")
    if not (TOP_MARGIN - MARGIN_TOLERANCE <= top_mm <= TOP_MARGIN + MARGIN_TOLERANCE):
        errors.append(f"Верхнее поле: {top_mm:.1f} мм (должно быть {TOP_MARGIN} мм)")
    if not (BOTTOM_MARGIN - MARGIN_TOLERANCE <= bottom_mm <= BOTTOM_MARGIN + MARGIN_TOLERANCE):
        errors.append(f"Нижнее поле: {bottom_mm:.1f} мм (должно быть {BOTTOM_MARGIN} мм)")
    return errors


def check_font(run, para):
    """Проверяет шрифт в run или параграфе, включая варианты Calibri и Calibri Light"""
    current_font = None
    if run.font.name and run.font.name != "None":
        current_font = run.font.name
    if not current_font:
        try:
            if para.style.font.name and para.style.font.name != "None":
                current_font = para.style.font.name
        except:
            pass

    if current_font and "calibri" in current_font.lower():
        return "Calibri"
    return current_font


def check_line_spacing(paragraph_format):
    """Проверяет, что межстрочный интервал равен 1.5"""
    if paragraph_format.line_spacing is None:
        return False
    if abs(paragraph_format.line_spacing - LINE_SPACING) < 0.01:
        return True
    if paragraph_format.line_spacing_rule in [4, 5]:
        line_spacing_pt = paragraph_format.line_spacing / Pt(1)
        expected_spacing_pt = MAIN_FONT_SIZE * LINE_SPACING
        return abs(line_spacing_pt - expected_spacing_pt) < 0.1
    return False


def check_paragraph_spacing(paragraph_format):
    """Проверяет, что интервалы перед и после абзаца равны 0 pt"""
    if (paragraph_format.space_before is not None and paragraph_format.space_before.pt > 0) or \
            (paragraph_format.space_after is not None and paragraph_format.space_after.pt > 0):
        return False
    return True


def check_first_line_indent(paragraph_format):
    """Проверяет отступ первой строки с учетом всех возможных способов его задания"""
    # Допустимое отклонение в см
    TOLERANCE_CM = 0.1

    # Если отступ явно не указан (None), считаем что проверка не требуется
    if (paragraph_format.first_line_indent is None and
            (not hasattr(paragraph_format, 'style') or
             (hasattr(paragraph_format.style, 'paragraph_format') and
              paragraph_format.style.paragraph_format.first_line_indent is None))):
        return True

    # Проверяем прямое задание отступа
    if paragraph_format.first_line_indent is not None:
        indent_cm = paragraph_format.first_line_indent.cm
        if abs(indent_cm - PARAGRAPH_INDENT_CM) <= TOLERANCE_CM:
            return True

    # Проверяем стиль абзаца
    if hasattr(paragraph_format, 'style') and hasattr(paragraph_format.style, 'paragraph_format'):
        style_indent = paragraph_format.style.paragraph_format.first_line_indent
        if style_indent is not None:
            indent_cm = style_indent.cm
            if abs(indent_cm - PARAGRAPH_INDENT_CM) <= TOLERANCE_CM:
                return True

    # Проверяем XML-разметку
    pPr = paragraph_format.element
    if pPr is not None:
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            first_line = ind.get(qn('w:firstLine'))
            if first_line is not None:
                try:
                    first_line_cm = Cm(int(first_line) / 1440).cm  # Конвертация twips в cm
                    if abs(first_line_cm - PARAGRAPH_INDENT_CM) <= TOLERANCE_CM:
                        return True
                except (ValueError, TypeError):
                    pass

    return False


def find_title_page_end(doc):
    """Определяет конец титульного листа (первый заголовок после титула)"""
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()
        if any(header in text for header in REQUIRED_HEADERS_LOWER + OPTIONAL_HEADERS_LOWER):
            return i
    return min(5, len(doc.paragraphs))  # Если заголовки не найдены, считаем первые 5 параграфов титульным листом


def fix_document(filename: str) -> str:
    """Исправляет ошибки в документе и возвращает путь к исправленному файлу"""
    doc = docx.Document(filename)

    # Определяем конец титульного листа
    title_page_end = find_title_page_end(doc)

    # Исправляем поля для всего документа
    for section in doc.sections:
        section.left_margin = Mm(LEFT_MARGIN)
        section.right_margin = Mm(RIGHT_MARGIN)
        section.top_margin = Mm(TOP_MARGIN)
        section.bottom_margin = Mm(BOTTOM_MARGIN)

    # Исправляем стили текста, пропуская титульный лист
    for para_idx, para in enumerate(doc.paragraphs):
        if para_idx < title_page_end:
            continue  # Пропускаем титульный лист

        # Устанавливаем шрифт Times New Roman и размер
        for run in para.runs:
            run.font.name = MAIN_FONT_NAME
            run.font.size = Pt(MAIN_FONT_SIZE)

        # Определяем тип параграфа
        style_name = para.style.name.lower()
        is_heading = style_name.startswith('heading')
        is_title = style_name in ['title', 'subtitle']
        is_header = is_heading or is_title
        text_lower = para.text.strip().lower()

        # Для заголовков устанавливаем жирный шрифт
        if is_header:
            for run in para.runs:
                run.font.bold = True

        # Устанавливаем выравнивание
        if is_header:
            # Для обязательных и необязательных заголовков - по центру
            if any(header in text_lower for header in REQUIRED_HEADERS_LOWER + OPTIONAL_HEADERS_LOWER):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Для остальных заголовков - по левому краю
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            # Для обычного текста - по ширине
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Устанавливаем межстрочный интервал 1.5
        para.paragraph_format.line_spacing = LINE_SPACING

        # Устанавливаем интервалы перед и после абзаца в 0 pt
        para.paragraph_format.space_before = Pt(PARAGRAPH_SPACING)
        para.paragraph_format.space_after = Pt(PARAGRAPH_SPACING)

        # Устанавливаем отступ первой строки (если это не заголовок и не элемент списка)
        if not is_header and not is_list_paragraph(para):
            para.paragraph_format.first_line_indent = Cm(PARAGRAPH_INDENT_CM)
            para.paragraph_format.left_indent = Cm(0)
            para.paragraph_format.right_indent = Cm(0)

    # Сохраняем исправленный документ
    fixed_filename = f"fixed_{os.path.basename(filename)}"
    doc.save(fixed_filename)
    return fixed_filename


async def check_document(filename: str) -> List[Tuple[str, str]]:
    """Проверяет документ на соответствие требованиям"""
    errors = []
    doc = docx.Document(filename)

    margin_errors = check_margins(doc.sections[0])
    if margin_errors:
        errors.append(("❌ Неправильные размеры полей", "\n".join(margin_errors)))

    first_content_header_pos = None
    found_headers = []
    content_page_end = None
    is_content_page = False
    appendix_page_start = None
    is_appendix_section = False
    references_page_start = None
    is_references_section = False

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text or text.startswith('\x0c'):
            continue

        style_name = para.style.name.lower()
        is_header = (style_name.startswith('heading') or style_name in ['title', 'subtitle'])

        if is_header:
            if not is_bold(para):
                example_text = text[:50] + ('...' if len(text) > 50 else '')
                errors.append(("❌ Заголовок должен быть жирным", f"Заголовок: '{example_text}'"))

            text_lower = text.lower()
            for header in REQUIRED_HEADERS_LOWER + OPTIONAL_HEADERS_LOWER:
                if header in text_lower:
                    found_headers.append(header)
                    if first_content_header_pos is None:
                        first_content_header_pos = i

                    if "содержание" in text_lower:
                        is_content_page = True
                    elif is_content_page:
                        content_page_end = i
                        is_content_page = False

                    if "приложение" in text_lower and not text_lower.startswith("приложение "):
                        is_appendix_section = True
                        appendix_page_start = i

                    if "список источников литературы" in text_lower or "список использованных источников" in text_lower:
                        is_references_section = True
                        references_page_start = i
                    break

    title_page_end = first_content_header_pos if first_content_header_pos is not None else min(5, len(doc.paragraphs))
    if content_page_end is None and is_content_page:
        content_page_end = len(doc.paragraphs)

    missing_headers = []
    for header in REQUIRED_HEADERS_LOWER:
        found = False
        for i in range(title_page_end, len(doc.paragraphs)):
            para = doc.paragraphs[i]
            text = para.text.strip().lower()
            if header in text:
                found = True
                break
        if not found:
            missing_headers.append(header)

    if missing_headers:
        errors.append(("❌ Отсутствуют обязательные разделы",
                       ", ".join(h.capitalize() for h in missing_headers)))

    for para_idx, para in enumerate(doc.paragraphs):
        if para_idx < title_page_end or (content_page_end and title_page_end <= para_idx < content_page_end):
            continue

        text = para.text.strip()
        if not text:
            continue

        if text.lower().startswith('приложение'):
            if not APPENDIX_PATTERN.match(text):
                break
            else:
                continue

        skip_indent_checks = is_references_section and para_idx >= references_page_start

        if not para.runs:
            continue

        if text.startswith('\x0c'):
            continue

        style_name = para.style.name.lower()
        is_heading = style_name.startswith('heading')
        is_title = style_name in ['title', 'subtitle']
        is_header = is_heading or is_title
        is_list = is_list_paragraph(para)

        para_errors = defaultdict(set)
        example_text = text[:50] + ('...' if len(text) > 50 else '')

        # Проверка интервалов между абзацами
        if not check_paragraph_spacing(para.paragraph_format):
            space_before = para.paragraph_format.space_before.pt if para.paragraph_format.space_before else 0
            space_after = para.paragraph_format.space_after.pt if para.paragraph_format.space_after else 0
            para_errors["❌ Интервалы перед и после абзаца должны быть 0 pt"].add(
                f"Пример: '{example_text}'\nТекущие значения: перед {space_before} pt, после {space_after} pt")

        if is_header:
            text_lower = text.lower()
            is_required_header = any(header in text_lower for header in REQUIRED_HEADERS_LOWER)
            is_content_or_appendix = any(header in text_lower for header in OPTIONAL_HEADERS_LOWER)
            is_left_aligned_header = (not is_required_header and not is_content_or_appendix)

            alignment = get_alignment(para)
            if is_required_header or is_content_or_appendix:
                if alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    current_alignment = "по левому краю" if alignment == WD_ALIGN_PARAGRAPH.LEFT else \
                        "по правому краю" if alignment == WD_ALIGN_PARAGRAPH.RIGHT else \
                            "по ширине" if alignment == WD_ALIGN_PARAGRAPH.JUSTIFY else \
                                "не определено"
                    para_errors["❌ Заголовок должен быть выровнен по центру"].add(
                        f"Пример: '{example_text}'\nТекущее выравнивание: {current_alignment}")
            else:
                if alignment != WD_ALIGN_PARAGRAPH.LEFT:
                    current_alignment = "по центру" if alignment == WD_ALIGN_PARAGRAPH.CENTER else \
                        "по правому краю" if alignment == WD_ALIGN_PARAGRAPH.RIGHT else \
                            "по ширине" if alignment == WD_ALIGN_PARAGRAPH.JUSTIFY else \
                                "не определено"
                    para_errors["❌ Заголовок должен быть выровнен по левому краю"].add(
                        f"Пример: '{example_text}'\nТекущее выравнивание: {current_alignment}")

            font_errors = set()
            size_errors = set()

            for run in para.runs:
                run_text = run.text.strip()
                if not run_text:
                    continue

                current_font = check_font(run, para)
                if current_font and current_font.lower() != MAIN_FONT_NAME.lower():
                    font_errors.add(current_font)

                current_size = run.font.size.pt if run.font.size else None
                if not current_size:
                    try:
                        current_size = para.style.font.size.pt
                    except:
                        pass

                if current_size and abs(current_size - MAIN_FONT_SIZE) > 0.1:
                    size_errors.add(f"{current_size:.1f}pt")

            if font_errors:
                para_errors[f"❌ Шрифт заголовка должен быть {MAIN_FONT_NAME}"].add(
                    f"Обнаружены шрифты: {', '.join(font_errors)}\nПример: '{example_text}'"
                )

            if size_errors:
                para_errors[f"❌ Размер шрифта заголовка должен быть {MAIN_FONT_SIZE}pt"].add(
                    f"Обнаружены размеры: {', '.join(size_errors)}\nПример: '{example_text}'"
                )

            if is_left_aligned_header:
                if not check_first_line_indent(para.paragraph_format):
                    para_errors[f"❌ Отступ первой строки должен быть {PARAGRAPH_INDENT_CM} см"].add(
                        f"Пример: '{example_text}'")

            for error_type, error_texts in para_errors.items():
                errors.append((error_type, "\n".join(error_texts)))
            continue

        if not is_list and not skip_indent_checks:
            left_indent = para.paragraph_format.left_indent
            right_indent = para.paragraph_format.right_indent

            # Проверка отступа первой строки
            if not check_first_line_indent(para.paragraph_format):
                para_errors[f"❌ Отступ первой строки должен быть {PARAGRAPH_INDENT_CM} см"].add(
                    f"Пример: '{example_text}'")

            if left_indent is not None and abs(left_indent.cm) > 0.1:
                para_errors["❌ Отступ слева должен быть 0 см"].add(
                    f"Текущее значение: {left_indent.cm:.1f} см\nПример: '{example_text}'")

            if right_indent is not None and abs(right_indent.cm) > 0.1:
                para_errors["❌ Отступ справа должен быть 0 см"].add(
                    f"Текущее значение: {right_indent.cm:.1f} см\nПример: '{example_text}'")

        font_errors = set()
        size_errors = set()

        for run in para.runs:
            run_text = run.text.strip()
            if not run_text:
                continue

            current_font = check_font(run, para)
            if current_font and current_font.lower() != MAIN_FONT_NAME.lower():
                font_errors.add(current_font)

            current_size = run.font.size.pt if run.font.size else None
            if not current_size:
                try:
                    current_size = para.style.font.size.pt
                except:
                    pass

            if current_size and abs(current_size - MAIN_FONT_SIZE) > 0.1:
                size_errors.add(f"{current_size:.1f}pt")

        if font_errors:
            para_errors[f"❌ Шрифт должен быть {MAIN_FONT_NAME}"].add(
                f"Обнаружены шрифты: {', '.join(font_errors)}\nПример: '{example_text}'"
            )

        if size_errors:
            para_errors[f"❌ Размер шрифта должен быть {MAIN_FONT_SIZE}pt"].add(
                f"Обнаружены размеры: {', '.join(size_errors)}\nПример: '{example_text}'"
            )

        if get_alignment(para) != WD_ALIGN_PARAGRAPH.JUSTIFY:
            para_errors["❌ Основной текст должен быть по ширине"].add(f"Пример: '{example_text}'")

        if not check_line_spacing(para.paragraph_format):
            para_errors[f"❌ Межстрочный интервал должен быть {LINE_SPACING}"].add(f"Пример: '{example_text}'")

        for error_type, error_texts in para_errors.items():
            errors.append((error_type, "\n".join(error_texts)))

    appendices = []
    invalid_appendix_letters = []

    for para_idx, para in enumerate(doc.paragraphs):
        if para_idx < title_page_end or (content_page_end and title_page_end <= para_idx < content_page_end):
            continue

        text = para.text.strip()
        if not text:
            continue

        if text.lower().startswith('приложение'):
            parts = text.split()
            if len(parts) == 2:
                letter = parts[1].upper()

                if letter in FORBIDDEN_LETTERS:
                    errors.append((f"❌ Запрещенное приложение с буквой {letter}", f"Текст: '{text[:50]}...'"))
                    continue

                if APPENDIX_PATTERN.match(text):
                    appendices.append((text, letter))
                else:
                    invalid_appendix_letters.append(letter)

    for letter in invalid_appendix_letters:
        errors.append(("❌ Недопустимая буква в приложении",
                       f"Использовано: '{letter}', допустимы: {', '.join(APPENDIX_ORDER[:5])}..."))

    if appendices:
        current_order = [letter for _, letter in appendices]

        for i, (appendix_text, current_letter) in enumerate(appendices):
            if i >= len(APPENDIX_ORDER):
                errors.append(("❌ Слишком много приложений",
                               f"Текст: '{appendix_text[:50]}...'\nМаксимально допустимое: Приложение {APPENDIX_ORDER[-1]}"))
                break

            expected_letter = APPENDIX_ORDER[i]
            if current_letter != expected_letter:
                errors.append((f"❌ Нарушен порядок приложений. Ожидалось: Приложение {expected_letter}",
                               f"Текст: '{appendix_text[:50]}...'"))
                break

        seen_letters = set()
        duplicates = set()
        for _, letter in appendices:
            if letter in seen_letters:
                duplicates.add(letter)
            seen_letters.add(letter)

        if duplicates:
            dup_messages = [f"Приложение {letter}" for letter in sorted(duplicates)]
            errors.append(("❌ Обнаружены дублирующиеся приложения", ", ".join(dup_messages)))

    return errors


@dp.message(CommandStart())
async def hello_start(message: types.Message):
    await message.answer(text=f'Добрый день, {message.from_user.full_name}! Пожалуйста, отправьте в чат вашу работу.')


@dp.message(lambda message: message.document)
async def handle_docx_file(message: types.Message):
    if not message.document.file_name.lower().endswith('.docx'):
        await message.reply("❌ Файл должен быть в формате .docx!")
        return

    await message.reply("✅ Файл принят! Обрабатываю...")

    try:
        file_id = message.document.file_id
        file = await bot.get_file(file_id)
        downloaded_file = await bot.download_file(file.file_path)

        temp_filename = f"temp_{message.from_user.id}.docx"
        with open(temp_filename, "wb") as f:
            f.write(downloaded_file.read())

        errors = await check_document(temp_filename)

        if errors:
            error_groups = defaultdict(list)
            for error, text in errors:
                error_groups[error].append(text)

            error_msg = "❌ Найдены ошибки в оформлении:\n\n"
            for error, texts in error_groups.items():
                error_msg += f"{error}:\n"
                unique_examples = set()
                for text in texts:
                    unique_examples.add(text.split("Пример:")[-1].strip())

                for example in list(unique_examples)[:3]:
                    error_msg += f"• {example}\n"

                if len(unique_examples) > 3:
                    error_msg += f"• и ещё {len(unique_examples) - 3} подобных случаев\n"
                error_msg += "\n"

            # Сохраняем временный файл для возможного исправления
            user_documents[message.from_user.id] = temp_filename

            # Создаем клавиатуру с кнопкой "Исправить"
            builder = InlineKeyboardBuilder()
            builder.add(types.InlineKeyboardButton(
                text="Исправить ошибки в документе",
                callback_data=f"fix_document_{message.from_user.id}")
            )

            await message.reply(error_msg[:4000], reply_markup=builder.as_markup())
        else:
            await message.reply("✅ Документ соответствует всем требованиям!")
            if os.path.exists(temp_filename):
                os.remove(temp_filename)

    except Exception as e:
        await message.reply(f"⚠ Ошибка при обработке файла: {str(e)}")
        if 'temp_filename' in locals() and os.path.exists(temp_filename):
            os.remove(temp_filename)


@dp.callback_query(lambda c: c.data.startswith('fix_document_'))
async def process_fix_document(callback: types.CallbackQuery):
    user_id = int(callback.data.split('_')[-1])

    if user_id not in user_documents:
        await callback.answer("Файл для исправления не найден. Пожалуйста, отправьте документ снова.")
        return

    temp_filename = user_documents[user_id]

    try:
        await callback.answer("Исправляю документ...")

        # Исправляем документ
        fixed_filename = fix_document(temp_filename)

        # Отправляем исправленный документ
        with open(fixed_filename, 'rb') as file:
            await bot.send_document(
                chat_id=callback.message.chat.id,
                document=types.BufferedInputFile(file.read(), filename=f"Исправленный_{fixed_filename}"),
                caption="✅ Вот исправленный документ"
            )

        # Удаляем временные файлы
        if os.path.exists(temp_filename):
            os.remove(temp_filename)
        if os.path.exists(fixed_filename):
            os.remove(fixed_filename)

        # Удаляем запись о файле
        del user_documents[user_id]

    except Exception as e:
        await callback.answer(f"Ошибка при исправлении документа: {str(e)}")
        if os.path.exists(temp_filename):
            os.remove(temp_filename)
        if 'fixed_filename' in locals() and os.path.exists(fixed_filename):
            os.remove(fixed_filename)


async def main():
    await dp.start_polling(bot)


if __name__ == "__main__":
    asyncio.run(main())